from docx import Document
from docx.shared import Pt
import os
from datetime import datetime


def create_word_document(content):

    # Create output folder
    os.makedirs("output", exist_ok=True)

    # Unique file name
    filename = f"AI_Document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    filepath = os.path.join("output", filename)

    # Create document
    document = Document()

    # Title
    title = document.add_heading("AI Generated Business Document", level=1)
    title.runs[0].font.size = Pt(20)

    # Date
    document.add_paragraph(
        f"Generated On: {datetime.now().strftime('%d-%m-%Y %H:%M:%S')}"
    )

    document.add_paragraph("-" * 60)

    # Split AI content into paragraphs
    paragraphs = content.split("\n")

    for para in paragraphs:

        para = para.strip()

        if para:

            if para.startswith("#"):

                document.add_heading(
                    para.replace("#", "").strip(),
                    level=2
                )

            elif para.startswith("*"):

                document.add_paragraph(
                    para.replace("*", "").strip(),
                    style="List Bullet"
                )

            else:

                document.add_paragraph(para)

    document.save(filepath)

    return filepath